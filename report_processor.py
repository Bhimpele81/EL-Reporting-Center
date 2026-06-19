"""
EL Reporting Center — Report Processor
----------------------------------------
Transforms raw camp management CSV exports into formatted Excel workbooks.
"""

import csv
import datetime
import io
import json
import os
import re
from datetime import date

from openpyxl import Workbook
from openpyxl.styles import (
    Alignment, Border, Font, PatternFill, Side
)
from openpyxl.utils import get_column_letter


# ---------------------------------------------------------------------------
# Config helpers
# ---------------------------------------------------------------------------

def load_bunk_config(config_path: str) -> dict:
    with open(config_path, "r") as f:
        return json.load(f)


def save_bunk_config(config_path: str, config: dict) -> None:
    with open(config_path, "w") as f:
        json.dump(config, f, indent=2)


def _norm(s: str) -> str:
    """Normalize a bunk name for matching: strip, collapse spaces, normalize dashes."""
    return re.sub(r"\s+", " ", str(s).strip()).replace("–", "-").replace("—", "-")


def get_bunk_lookup(config: dict) -> dict:
    """Return {normalized_bunk_name: {number, camp, original}} from config."""
    lookup = {}
    for camp in config.get("camps", []):
        for bunk in camp.get("bunks", []):
            lookup[_norm(bunk["name"])] = {
                "number": bunk["number"],
                "camp": camp["name"],
                "original": bunk["name"],
            }
    return lookup


def get_ordered_bunks(config: dict) -> list:
    """Return list of bunk names in display order (camp order, then bunk number order)."""
    bunks = []
    for camp in config.get("camps", []):
        for bunk in sorted(camp.get("bunks", []), key=lambda b: int(b.get("number") or 999)):
            bunks.append(_norm(bunk["name"]))
    return bunks


def _bunk_num(bunk_name) -> int:
    """Extract leading integer from a bunk name for numeric sort; non-numeric bunks sort last."""
    m = re.match(r'^(\d+)', str(bunk_name).strip())
    return int(m.group(1)) if m else 9999


def _bunk_sort_key(bunk_name):
    """Group/print order for bunks:
       0) numbered bunks (incl. PT CITs) by number
       1) FT CITs (after PT CITs)
       2) any other named bunk
       3) Staff Transport, then unassigned (blank) — always last.
    """
    b  = str(bunk_name).strip()
    bl = b.lower()
    m  = re.match(r'^(\d+)', b)
    if m:
        return (0, int(m.group(1)), bl)
    if "ft cit" in bl:
        return (1, 0, bl)
    if not b:
        return (3, 1, bl)            # unassigned / blank bunk — very last
    if "staff transport" in bl:
        return (3, 0, bl)            # Staff Transport — last group
    return (2, 0, bl)


# ---------------------------------------------------------------------------
# Grade normalizer
# ---------------------------------------------------------------------------

def normalize_grade(raw: str) -> str:
    g = str(raw).strip()
    if not g or g.lower() == "nan":
        return ""
    gl = g.lower()
    if "pre-k" in gl or gl in ("prek", "pre k", "pk"):
        return "PK"
    if gl.startswith("pre"):          # Pre-School / Preschool
        return "PS"
    if gl in ("k", "kindergarten", "kg"):
        return "K"
    m = re.match(r"^(\d+)", g)
    if m:
        return m.group(1)
    return g


# ---------------------------------------------------------------------------
# Raw CSV parser
# ---------------------------------------------------------------------------

WEEK_RE = re.compile(r"Week\s+(\d+)", re.IGNORECASE)


def _detect_col(header: list, keywords: list, fallback: int) -> int:
    """Return the first column index whose header contains ALL keywords (case-insensitive)."""
    for i, h in enumerate(header):
        h_lower = str(h).lower()
        if all(kw in h_lower for kw in keywords):
            return i
    return fallback


def _rows_to_campers(rows: list) -> list:
    """
    Convert a list of rows (list-of-strings) into camper dicts.

    Expected columns (0-indexed):
      0  row#
      1  Last name
      2  First name
      3  Bunk name   (e.g. "01 Munchkins")
      4  Session name (e.g. "Week 1, Week 3 (Camp Photos), Week 4")
      5  Age + months
      6  Current grade
      7  Monday?     (Yes / No / blank)
      8  Tuesday?
      9  Wednesday?
      10 Thursday?
      11 Friday?
      ?  Driver      (detected by header name — column position varies)
    """
    # Detect the driver column by scanning header for any cell containing "driver"
    header = rows[0] if rows else []
    driver_col = None
    for i, h in enumerate(header):
        if "driver" in str(h).strip().lower():
            driver_col = i
            break

    campers = []
    for row in rows[1:]:          # skip header
        if len(row) < 4 or not str(row[0]).strip().isdigit():
            continue

        last     = str(row[1]).strip()
        first    = str(row[2]).strip()
        bunk     = _norm(row[3])
        sessions = str(row[4]).strip() if len(row) > 4 else ""
        age      = str(row[5]).strip() if len(row) > 5 else ""
        grade    = normalize_grade(row[6]) if len(row) > 6 else ""
        mon      = str(row[7]).strip()  if len(row) > 7  else ""
        tue      = str(row[8]).strip()  if len(row) > 8  else ""
        wed      = str(row[9]).strip()  if len(row) > 9  else ""
        thu      = str(row[10]).strip() if len(row) > 10 else ""
        fri      = str(row[11]).strip() if len(row) > 11 else ""

        # Driver: read from detected column; skip None/nan/empty values
        raw_driver = str(row[driver_col]).strip() if (driver_col is not None and driver_col < len(row)) else ""
        driver = "" if raw_driver.lower() in ("", "none", "nan", "n/a", "#n/a") else raw_driver

        weeks = [0] * 8
        for part in sessions.split(","):
            m = WEEK_RE.search(part)
            if m:
                wk = int(m.group(1))
                if 1 <= wk <= 8:
                    weeks[wk - 1] = 1

        any_day_specified = any(
            d.lower() in ("yes", "no") for d in [mon, tue, wed, thu, fri]
        )
        if any_day_specified:
            day_m = "M" if mon.lower() == "yes" else None
            day_t = "T" if tue.lower() == "yes" else None
            day_w = "W" if wed.lower() == "yes" else None
            day_r = "R" if thu.lower() == "yes" else None
            day_f = "F" if fri.lower() == "yes" else None
        else:
            day_m, day_t, day_w, day_r, day_f = "M", "T", "W", "R", "F"

        campers.append({
            "name":   f"{last}, {first}",
            "bunk":   bunk,
            "weeks":  weeks,
            "days":   [day_m, day_t, day_w, day_r, day_f],
            "age":    age,
            "grade":  grade,
            "driver": driver,
        })

    return campers


def parse_raw_csv(file_bytes: bytes) -> list:
    """Parse a raw bunk-snapshot export — accepts CSV or XLSX."""
    # XLSX files start with the ZIP magic bytes PK\x03\x04
    if file_bytes[:4] == b'PK\x03\x04':
        from openpyxl import load_workbook
        wb = load_workbook(filename=io.BytesIO(file_bytes), read_only=True, data_only=True)
        ws = wb.active
        rows = [[str(cell.value) if cell.value is not None else "" for cell in row]
                for row in ws.iter_rows()]
        wb.close()
        return _rows_to_campers(rows)

    content = file_bytes.decode("utf-8-sig", errors="replace")
    reader  = csv.reader(io.StringIO(content))
    rows    = list(reader)
    return _rows_to_campers(rows)


def parse_driver_csv(file_bytes: bytes) -> list:
    """
    Parse a raw Driver Totals export (CSV or XLSX).

    Expected columns (0-indexed, detected by header name with fallbacks):
      0  row#
      1  Last name
      2  First name
      3  Driver        (header contains "driver")
      4  Bunk name     (header contains "bunk")
      5  Session name  (header contains "session")
      6  Age           (header contains "age")
      7  Grade         (header contains "grade")
      8  Monday?       (header contains "monday" or "mon")
      9  Tuesday?
      10 Wednesday?
      11 Thursday?
      12 Friday?
    """
    if file_bytes[:4] == b'PK\x03\x04':
        from openpyxl import load_workbook as _lw
        _wb = _lw(filename=io.BytesIO(file_bytes), read_only=True, data_only=True)
        _ws = _wb.active
        rows = [[str(c.value) if c.value is not None else "" for c in r]
                for r in _ws.iter_rows()]
        _wb.close()
    else:
        content = file_bytes.decode("utf-8-sig", errors="replace")
        rows = list(csv.reader(io.StringIO(content)))

    if not rows:
        return []

    header = rows[0]
    driver_col  = _detect_col(header, ["driver"],    3)
    stop_col    = _detect_col(header, ["stop"],      None)   # optional column
    bunk_col    = _detect_col(header, ["bunk"],      5)
    session_col = _detect_col(header, ["session"],   6)
    age_col     = _detect_col(header, ["age"],       7)
    grade_col   = _detect_col(header, ["grade"],     8)
    mon_col     = _detect_col(header, ["monday"],    9)
    tue_col     = _detect_col(header, ["tuesday"],   10)
    wed_col     = _detect_col(header, ["wednesday"], 11)
    thu_col     = _detect_col(header, ["thursday"],  12)
    fri_col     = _detect_col(header, ["friday"],    13)

    def _val(row, col):
        return str(row[col]).strip() if col < len(row) else ""

    campers = []
    for row in rows[1:]:
        if len(row) < 2 or not str(row[0]).strip().isdigit():
            continue

        last     = _val(row, 1)
        first    = _val(row, 2)
        raw_drv  = _val(row, driver_col)
        driver   = "" if raw_drv.lower() in ("", "none", "nan", "n/a", "#n/a") else raw_drv

        # Stop #: optional — only present in newer exports
        if stop_col is not None and stop_col < len(row):
            try:
                stop_val = int(float(_val(row, stop_col)))
            except (ValueError, TypeError):
                stop_val = None
        else:
            stop_val = None

        bunk     = _val(row, bunk_col)
        sessions = _val(row, session_col)
        raw_age  = _val(row, age_col)
        raw_grade = _val(row, grade_col)

        # Age: keep as float if possible
        try:
            age_val = float(raw_age)
        except (ValueError, TypeError):
            age_val = raw_age if raw_age else None

        grade_val = normalize_grade(raw_grade)

        mon = _val(row, mon_col)
        tue = _val(row, tue_col)
        wed = _val(row, wed_col)
        thu = _val(row, thu_col)
        fri = _val(row, fri_col)

        weeks = [0] * 8
        for part in sessions.split(","):
            m = WEEK_RE.search(part)
            if m:
                wk = int(m.group(1))
                if 1 <= wk <= 8:
                    weeks[wk - 1] = 1

        any_day_specified = any(d.lower() in ("yes", "no") for d in [mon, tue, wed, thu, fri])
        if any_day_specified:
            day_m = "M" if mon.lower() == "yes" else None
            day_t = "T" if tue.lower() == "yes" else None
            day_w = "W" if wed.lower() == "yes" else None
            day_r = "R" if thu.lower() == "yes" else None
            day_f = "F" if fri.lower() == "yes" else None
        else:
            day_m, day_t, day_w, day_r, day_f = "M", "T", "W", "R", "F"

        campers.append({
            "name":   f"{last}, {first}",
            "bunk":   bunk,
            "driver": driver,
            "stop":   stop_val,
            "weeks":  weeks,
            "days":   [day_m, day_t, day_w, day_r, day_f],
            "age":    age_val,
            "grade":  grade_val,
        })

    return campers


# ---------------------------------------------------------------------------
# Report builder
# ---------------------------------------------------------------------------

# ---- Styles ----------------------------------------------------------------

BRAND     = "6D1F2F"
BRAND_ALT = "F5E6E9"
WHITE     = "FFFFFF"
LIGHT_GREY = "F2F2F2"
DARK_GREY  = "1A1018"

_thin = Side(style="thin", color="CCCCCC")
_med  = Side(style="medium", color="AAAAAA")
THIN_BORDER = Border(left=_thin, right=_thin, top=_thin, bottom=_thin)
MED_BORDER  = Border(left=_med,  right=_med,  top=_med,  bottom=_med)

HEADER_FONT   = Font(name="Calibri", bold=True, color=WHITE, size=10)
SUBHDR_FONT   = Font(name="Calibri", bold=True, size=10)
BODY_FONT     = Font(name="Calibri", size=10)
TOTAL_FONT    = Font(name="Calibri", bold=True, size=10)
DATE_FONT     = Font(name="Calibri", bold=True, size=11)

BRAND_FILL    = PatternFill("solid", fgColor=BRAND)
ALT_FILL      = PatternFill("solid", fgColor=BRAND_ALT)
LGREY_FILL    = PatternFill("solid", fgColor="EEEEEE")
TOTAL_FILL    = PatternFill("solid", fgColor="D9D9D9")

CENTER = Alignment(horizontal="center", vertical="center")
LEFT   = Alignment(horizontal="left",   vertical="center")
RIGHT  = Alignment(horizontal="right",  vertical="center")


def _cell(ws, row, col, value, font=None, fill=None, align=None, border=None):
    c = ws.cell(row=row, column=col, value=value)
    if font:   c.font   = font
    if fill:   c.fill   = fill
    if align:  c.alignment = align
    if border: c.border = border
    return c


# ---------------------------------------------------------------------------
# Build the "Report" sheet
# ---------------------------------------------------------------------------

def build_report_sheet(ws, campers: list, bunk_lookup: dict,
                        ordered_bunks: list, report_date: date):

    # ----- Local (larger) fonts for the Report sheet ------------------------
    # Defined locally so they don't change the Totals sheet's shared styles.
    R_HEADER = Font(name="Calibri", bold=True, color=WHITE, size=11)
    R_BODY   = Font(name="Calibri", size=13)
    R_TOTAL  = Font(name="Calibri", bold=True, size=13)
    R_BUNK   = Font(name="Calibri", bold=True, color="000000", size=20)  # black bunk name

    # Dark vertical separator (medium weight) drawn on a cell's left edge
    SEP_LEFT = Border(left=Side(style="medium"))

    # ----- Column layout (Bunk column removed) ------------------------------
    #   A(1)=Child   B-I(2-9)=#1-#8   J-N(10-14)=Days M T W R F
    #   O(15)=Age    P(16)=Grade
    # The bunk name prints as a large black title at the top-left of each
    # bunk's block (each bunk starts on its own page). The report date is in
    # the page footer.
    COL_CHILD = 1
    COL_WK1   = 2     # weeks occupy cols 2-9
    COL_DAY1  = 10    # days  occupy cols 10-14
    COL_AGE   = 15
    COL_GRADE = 16
    LAST_COL  = 16

    def _write_col_headers(hr):
        ws.row_dimensions[hr].height = 16
        # Child + weeks: maroon fill, no borders
        for ci, h in [(COL_CHILD, "Child"),
                      (2, "#1"), (3, "#2"), (4, "#3"), (5, "#4"),
                      (6, "#5"), (7, "#6"), (8, "#7"), (9, "#8")]:
            c = ws.cell(row=hr, column=ci, value=h)
            c.font = R_HEADER; c.fill = BRAND_FILL; c.alignment = CENTER
        # Days header (merged) — dark separator on its left, no other borders
        ws.merge_cells(start_row=hr, start_column=COL_DAY1,
                       end_row=hr,   end_column=COL_DAY1 + 4)
        dcell = ws.cell(row=hr, column=COL_DAY1, value="Days")
        dcell.font = R_HEADER; dcell.fill = BRAND_FILL
        dcell.alignment = CENTER; dcell.border = SEP_LEFT
        for di in range(1, 5):
            ws.cell(row=hr, column=COL_DAY1 + di).fill = BRAND_FILL
        # Age / Grade headers — dark separator on each left
        ac = ws.cell(row=hr, column=COL_AGE, value="Age")
        ac.font = R_HEADER; ac.fill = BRAND_FILL; ac.alignment = CENTER; ac.border = SEP_LEFT
        gc = ws.cell(row=hr, column=COL_GRADE, value="Grade")
        gc.font = R_HEADER; gc.fill = BRAND_FILL; gc.alignment = CENTER; gc.border = SEP_LEFT

    # ----- Group campers by bunk -------------------------------------------
    bunk_groups = {}
    for c in campers:
        bunk_groups.setdefault(c["bunk"], []).append(c)

    # Sort campers alphabetically within each bunk
    for bk in bunk_groups:
        bunk_groups[bk].sort(key=lambda x: x["name"])

    display_order = sorted(bunk_groups.keys(), key=_bunk_sort_key)

    # ----- Write rows -------------------------------------------------------
    # No global date row (date is in the footer); bunk blocks start at row 1.
    from openpyxl.worksheet.pagebreak import Break
    row = 1
    max_a_len = len("Total:   00")   # track widest column-A value for autofit

    for bk_idx, bunk_name in enumerate(display_order):
        group = bunk_groups[bunk_name]
        week_sums = [0] * 8

        # --- Bunk name (large black title, top-left, no fill/border) ---
        bt = ws.cell(row=row, column=1, value=bunk_name)
        bt.font = R_BUNK; bt.alignment = LEFT
        ws.row_dimensions[row].height = 26
        row += 1

        # --- Column headers (under each bunk name) ---
        _write_col_headers(row)
        row += 1

        for ci, camper in enumerate(group):
            ws.row_dimensions[row].height = 18
            alt = (ci % 2 == 1)
            fill = ALT_FILL if alt else None

            # Child + weeks (A-I): no gridlines (alt shading carries the rows)
            _cell(ws, row, COL_CHILD, camper["name"], font=R_BODY, fill=fill, align=LEFT)
            max_a_len = max(max_a_len, len(str(camper["name"] or "")))

            for wi, wv in enumerate(camper["weeks"]):
                _cell(ws, row, COL_WK1 + wi, wv,
                      font=R_BODY, fill=fill, align=CENTER)
                week_sums[wi] += wv

            # Days (J-N): no gridlines, except dark separator left of the block
            for di, dv in enumerate(camper["days"]):
                _cell(ws, row, COL_DAY1 + di, dv, font=R_BODY, fill=fill,
                      align=CENTER, border=(SEP_LEFT if di == 0 else None))

            # Age as a number; grade as a number when numeric, else keep the
            # normalized text (K / PK / PS) so non-numeric grades still show.
            age_val = camper["age"]
            try: age_val = float(str(age_val).strip())
            except (ValueError, TypeError): age_val = None
            gtext = str(camper["grade"]).strip()
            try:
                grade_val = int(gtext)
            except (ValueError, TypeError):
                grade_val = gtext or None

            # Age / Grade: dark separator on each left edge
            _cell(ws, row, COL_AGE,   age_val,   font=R_BODY, fill=fill, align=CENTER, border=SEP_LEFT)
            _cell(ws, row, COL_GRADE, grade_val, font=R_BODY, fill=fill, align=CENTER, border=SEP_LEFT)
            row += 1

        # --- Subtotal row: total count under the names in column A ---
        ws.row_dimensions[row].height = 18
        _cell(ws, row, COL_CHILD, f"Total:   {len(group)}",
              font=R_TOTAL, fill=TOTAL_FILL, align=LEFT)
        # Week-sum cells (B-I): no gridlines, fill only
        for wi, wsum in enumerate(week_sums):
            _cell(ws, row, COL_WK1 + wi, wsum,
                  font=R_TOTAL, fill=TOTAL_FILL, align=CENTER)
        # Days (J-N): no gridlines, dark separator left of the block
        for di in range(5):
            _cell(ws, row, COL_DAY1 + di, None, fill=TOTAL_FILL,
                  border=(SEP_LEFT if di == 0 else None))
        _cell(ws, row, COL_AGE,   None, fill=TOTAL_FILL, border=SEP_LEFT)
        _cell(ws, row, COL_GRADE, None, fill=TOTAL_FILL, border=SEP_LEFT)
        row += 1

        # Page break after each bunk (except the last)
        if bk_idx < len(display_order) - 1:
            ws.row_breaks.append(Break(id=row - 1))

    # ----- Column widths (wider to use the freed-up space) ------------------
    last_row = row - 1
    ws.column_dimensions["A"].width = max(20, int(max_a_len * 1.15))
    for col_letter in [get_column_letter(c) for c in range(COL_WK1, COL_WK1 + 8)]:
        ws.column_dimensions[col_letter].width = 8   # #1-#8
    for col_letter in [get_column_letter(c) for c in range(COL_DAY1, COL_DAY1 + 5)]:
        ws.column_dimensions[col_letter].width = 3.5   # M T W R F (tight together)
    ws.column_dimensions[get_column_letter(COL_AGE)].width   = 8
    ws.column_dimensions[get_column_letter(COL_GRADE)].width = 8

    # ----- Suppress green error indicators in Age/Grade --------------------
    # openpyxl 3.1.x has no IgnoredErrors helper — inject XML directly
    if last_row >= 2:
        try:
            from lxml import etree
            ns = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
            ie_elem = etree.SubElement(ws._element,
                                       f"{{{ns}}}ignoredErrors")
            err = etree.SubElement(ie_elem, f"{{{ns}}}ignoredError")
            err.set("sqref",
                    f"{get_column_letter(COL_AGE)}2:{get_column_letter(COL_GRADE)}{last_row}")
            err.set("numberStoredAsText", "1")
            err.set("formula", "1")
            err.set("formulaRange", "1")
        except Exception:
            pass

    # Hide default gridlines; structure comes from the drawn borders, leaving
    # the week and day columns (which have no borders) clean.
    ws.sheet_view.showGridLines = False

    # ----- Print settings: landscape, fit to 1 page wide --------------------
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToPage   = True
    ws.page_setup.fitToWidth  = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.print_title_rows = None

    # ----- Margins (inches) -------------------------------------------------
    ws.page_margins.left   = 0.25
    ws.page_margins.right  = 0.25
    ws.page_margins.top    = 0.5
    ws.page_margins.bottom = 0.5
    ws.page_margins.footer = 0.25   # footer sits 0.25" above the bottom edge

    # Center the table horizontally so left/right margins look equal
    ws.print_options.horizontalCentered = True

    # ----- Footer: report date on every page --------------------------------
    date_str = (report_date.strftime("%-m/%-d/%Y") if os.name != "nt"
                else report_date.strftime("%#m/%#d/%Y"))
    ws.oddFooter.right.text  = f"&12Report Date: {date_str}"
    ws.evenFooter.right.text = f"&12Report Date: {date_str}"


# ---------------------------------------------------------------------------
# Build the "Totals" sheet
# ---------------------------------------------------------------------------

def build_totals_sheet(ws, campers: list, config: dict,
                        bunk_lookup: dict, report_date: date):

    # Pre-compute per-bunk counts and week totals
    bunk_count = {}   # bunk_name -> total campers
    bunk_weeks = {}   # bunk_name -> [w1..w8]

    for c in campers:
        bk = c["bunk"]
        bunk_count[bk] = bunk_count.get(bk, 0) + 1
        if bk not in bunk_weeks:
            bunk_weeks[bk] = [0] * 8
        for wi, wv in enumerate(c["weeks"]):
            bunk_weeks[bk][wi] += wv

    # Per-camp totals
    camp_count = {}   # camp -> total
    camp_weeks = {}   # camp -> [w1..w8]
    for camp in config["camps"]:
        cn = camp["name"]
        camp_count[cn] = 0
        camp_weeks[cn] = [0] * 8
        for bunk in camp["bunks"]:
            bk = bunk["name"]
            camp_count[cn] += bunk_count.get(bk, 0)
            for wi in range(8):
                camp_weeks[cn][wi] += bunk_weeks.get(bk, [0]*8)[wi]

    grand_total = sum(camp_count.values())
    grand_weeks = [sum(camp_weeks[c][wi] for c in camp_weeks) for wi in range(8)]

    # ---- Larger local fonts + layout --------------------------------------
    T_HDR   = Font(name="Calibri", bold=True, color=WHITE, size=14)
    T_SUB   = Font(name="Calibri", bold=True, size=13)
    T_BODY  = Font(name="Calibri", size=13)
    T_TOTAL = Font(name="Calibri", bold=True, size=13)
    T_DATE  = Font(name="Calibri", bold=True, size=12)
    ROW_H   = 17

    #  LEFT column  : A=Camp  B=Bunk  C=Total  (Bunk Totals, then Group Totals below)
    #  GAP          : col D
    #  RIGHT column : E=Label  F-M=#1..#8  (Group Totals by Week, then Bunk Totals by Week)
    L_CAMP, L_BUNK, L_TOT = 1, 2, 3
    R_LABEL = 5
    R_W1    = 6
    R_END   = R_W1 + 7   # column M

    camp_names = [c["name"] for c in config["camps"]]

    def _grp_label(cn):
        # The Upper total excludes full-time CITs (their own group) — note it
        return f"{cn} (w/o FT)" if cn.strip().lower().startswith("upper") else cn

    all_bunks_ordered, bunk_camp = [], {}
    for camp in config["camps"]:
        for b in sorted(camp["bunks"], key=lambda b: int(b.get("number") or 999)):
            all_bunks_ordered.append(b["name"])
            bunk_camp[b["name"]] = camp["name"]

    def _section_hdr(r, c1, c2, text):
        ws.merge_cells(start_row=r, start_column=c1, end_row=r, end_column=c2)
        cc = ws.cell(row=r, column=c1, value=text)
        cc.font = T_HDR; cc.fill = BRAND_FILL; cc.alignment = CENTER; cc.border = THIN_BORDER
        for c in range(c1 + 1, c2 + 1):
            ws.cell(row=r, column=c).fill = BRAND_FILL
            ws.cell(row=r, column=c).border = THIN_BORDER

    def _week_subhdr(r):
        _cell(ws, r, R_LABEL, None, fill=LGREY_FILL, border=THIN_BORDER)
        for wi in range(8):
            _cell(ws, r, R_W1 + wi, f"#{wi+1}",
                  font=T_SUB, fill=LGREY_FILL, align=CENTER, border=THIN_BORDER)

    # ----- Row 1: date ------------------------------------------------------
    _cell(ws, 1, 1, "Report Date", font=T_DATE)
    _cell(ws, 1, 2, report_date.strftime("%-m/%-d/%Y") if os.name != "nt"
          else report_date.strftime("%#m/%#d/%Y"), font=T_DATE)

    # ===== LEFT COLUMN: Bunk Totals =====
    _section_hdr(2, L_CAMP, L_TOT, "Bunk Totals")
    for ci, h in [(L_CAMP, "Camp"), (L_BUNK, "Bunk"), (L_TOT, "Total")]:
        _cell(ws, 3, ci, h, font=T_SUB, fill=LGREY_FILL, align=CENTER, border=THIN_BORDER)
    lr = 4
    for bi, bk in enumerate(all_bunks_ordered):
        fill = ALT_FILL if bi % 2 else None
        _cell(ws, lr, L_CAMP, bunk_camp.get(bk, ""), font=T_BODY, fill=fill, align=LEFT, border=THIN_BORDER)
        _cell(ws, lr, L_BUNK, bk, font=T_BODY, fill=fill, align=LEFT, border=THIN_BORDER)
        _cell(ws, lr, L_TOT, bunk_count.get(bk, 0), font=T_BODY, fill=fill, align=CENTER, border=THIN_BORDER)
        lr += 1
    _cell(ws, lr, L_CAMP, "TOTAL", font=T_TOTAL, fill=TOTAL_FILL, align=LEFT, border=THIN_BORDER)
    _cell(ws, lr, L_BUNK, None,    font=T_TOTAL, fill=TOTAL_FILL, border=THIN_BORDER)
    _cell(ws, lr, L_TOT, grand_total, font=T_TOTAL, fill=TOTAL_FILL, align=CENTER, border=THIN_BORDER)
    lr += 2   # gap row

    # ===== LEFT COLUMN: Group Totals (under Bunk Totals) =====
    _section_hdr(lr, L_CAMP, L_TOT, "Group Totals"); lr += 1
    ws.merge_cells(start_row=lr, start_column=L_CAMP, end_row=lr, end_column=L_BUNK)
    _cell(ws, lr, L_CAMP, "Camp",  font=T_SUB, fill=LGREY_FILL, align=CENTER, border=THIN_BORDER)
    _cell(ws, lr, L_BUNK, None,    fill=LGREY_FILL, border=THIN_BORDER)
    _cell(ws, lr, L_TOT,  "Total", font=T_SUB, fill=LGREY_FILL, align=CENTER, border=THIN_BORDER)
    lr += 1
    for ci, cn in enumerate(camp_names):
        fill = ALT_FILL if ci % 2 else None
        ws.merge_cells(start_row=lr, start_column=L_CAMP, end_row=lr, end_column=L_BUNK)
        _cell(ws, lr, L_CAMP, _grp_label(cn), font=T_BODY, fill=fill, align=LEFT, border=THIN_BORDER)
        _cell(ws, lr, L_BUNK, None, fill=fill, border=THIN_BORDER)
        _cell(ws, lr, L_TOT, camp_count[cn], font=T_BODY, fill=fill, align=CENTER, border=THIN_BORDER)
        lr += 1
    ws.merge_cells(start_row=lr, start_column=L_CAMP, end_row=lr, end_column=L_BUNK)
    _cell(ws, lr, L_CAMP, "Total", font=T_TOTAL, fill=TOTAL_FILL, align=LEFT, border=THIN_BORDER)
    _cell(ws, lr, L_BUNK, None,    fill=TOTAL_FILL, border=THIN_BORDER)
    _cell(ws, lr, L_TOT, grand_total, font=T_TOTAL, fill=TOTAL_FILL, align=CENTER, border=THIN_BORDER)
    left_last = lr

    # ===== RIGHT COLUMN: Group Totals by Week =====
    _section_hdr(2, R_LABEL, R_END, "Group Totals by Week")
    _week_subhdr(3)
    rr = 4
    for ci, cn in enumerate(camp_names):
        fill = ALT_FILL if ci % 2 else None
        _cell(ws, rr, R_LABEL, _grp_label(cn), font=T_BODY, fill=fill, align=LEFT, border=THIN_BORDER)
        for wi in range(8):
            _cell(ws, rr, R_W1 + wi, camp_weeks[cn][wi], font=T_BODY, fill=fill, align=CENTER, border=THIN_BORDER)
        rr += 1
    _cell(ws, rr, R_LABEL, "Total", font=T_TOTAL, fill=TOTAL_FILL, align=LEFT, border=THIN_BORDER)
    for wi in range(8):
        _cell(ws, rr, R_W1 + wi, grand_weeks[wi], font=T_TOTAL, fill=TOTAL_FILL, align=CENTER, border=THIN_BORDER)
    rr += 2   # gap row

    # ===== RIGHT COLUMN: Bunk Totals by Week =====
    _section_hdr(rr, R_LABEL, R_END, "Bunk Totals by Week"); rr += 1
    _week_subhdr(rr); rr += 1
    for bi, bk in enumerate(all_bunks_ordered):
        if bk not in bunk_weeks:
            continue
        fill = ALT_FILL if bi % 2 else None
        _cell(ws, rr, R_LABEL, bk, font=T_BODY, fill=fill, align=LEFT, border=THIN_BORDER)
        for wi in range(8):
            _cell(ws, rr, R_W1 + wi, bunk_weeks[bk][wi], font=T_BODY, fill=fill, align=CENTER, border=THIN_BORDER)
        rr += 1
    _cell(ws, rr, R_LABEL, "Total", font=T_TOTAL, fill=TOTAL_FILL, align=LEFT, border=THIN_BORDER)
    for wi in range(8):
        _cell(ws, rr, R_W1 + wi, grand_weeks[wi], font=T_TOTAL, fill=TOTAL_FILL, align=CENTER, border=THIN_BORDER)

    # ----- Row heights (taller) & column widths (wider, fill the page) ------
    for row in range(1, max(left_last, rr) + 1):
        ws.row_dimensions[row].height = ROW_H

    # Sized to fill a portrait (tall) page top-to-bottom and edge-to-edge
    ws.column_dimensions["A"].width = 10   # Camp
    ws.column_dimensions["B"].width = 19   # Bunk
    ws.column_dimensions["C"].width = 8    # Total
    ws.column_dimensions["D"].width = 2    # gap
    ws.column_dimensions[get_column_letter(R_LABEL)].width = 19   # Label
    for wi in range(8):
        ws.column_dimensions[get_column_letter(R_W1 + wi)].width = 6

    # ---- Print settings: portrait (tall), scaled to a SINGLE page ----------
    ws.page_setup.orientation = "portrait"
    ws.page_setup.fitToPage   = True
    ws.page_setup.fitToWidth  = 1
    ws.page_setup.fitToHeight = 1
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.print_options.horizontalCentered = True
    ws.print_options.verticalCentered   = True
    ws.page_margins.left   = 0.2
    ws.page_margins.right  = 0.2
    ws.page_margins.top    = 0.2
    ws.page_margins.bottom = 0.2


# ---------------------------------------------------------------------------
# Group Attendance parser + builder
# ---------------------------------------------------------------------------

def parse_group_attendance(file_bytes: bytes) -> list:
    """
    Parse raw group attendance CSV/XLSX export.

    Expected columns (0-indexed):
      0  row#
      1  Bunk name
      2  Last name
      3  First name
      4  Monday?   (Yes / No / blank)
      5  Tuesday?
      6  Wednesday?
      7  Thursday?
      8  Friday?
    """
    if file_bytes[:4] == b'PK\x03\x04':
        from openpyxl import load_workbook as _lw
        _wb = _lw(filename=io.BytesIO(file_bytes), read_only=True, data_only=True)
        _ws = _wb.active
        rows = [[str(c.value) if c.value is not None else "" for c in r]
                for r in _ws.iter_rows()]
        _wb.close()
    else:
        content = file_bytes.decode("utf-8-sig", errors="replace")
        rows = list(csv.reader(io.StringIO(content)))

    campers = []
    for row in rows[1:]:
        if len(row) < 4 or not str(row[0]).strip().isdigit():
            continue
        bunk  = str(row[1]).strip()
        last  = str(row[2]).strip()
        first = str(row[3]).strip()
        mon   = str(row[4]).strip() if len(row) > 4 else ""
        tue   = str(row[5]).strip() if len(row) > 5 else ""
        wed   = str(row[6]).strip() if len(row) > 6 else ""
        thu   = str(row[7]).strip() if len(row) > 7 else ""
        fri   = str(row[8]).strip() if len(row) > 8 else ""

        any_specified = any(d.lower() in ("yes", "no") for d in [mon, tue, wed, thu, fri])
        if any_specified:
            enrolled = (
                ("M" if mon.lower() == "yes" else "") +
                ("T" if tue.lower() == "yes" else "") +
                ("W" if wed.lower() == "yes" else "") +
                ("R" if thu.lower() == "yes" else "") +
                ("F" if fri.lower() == "yes" else "")
            )
            if enrolled == "MTWRF":
                enrolled = ""   # full week — treat same as blank
        else:
            enrolled = ""

        campers.append({"name": f"{last}, {first}", "bunk": bunk, "enrolled": enrolled})

    return campers


def build_group_attendance_sheet(ws, campers: list, config: dict,
                                  report_date=None, week_num: int = None) -> None:
    """
    Build the Data1 sheet for Group Attendance.

    Column layout:
      A  – Bunk name  (merged + rotated 90° for entire bunk group)
      B  – Camper     (bold 16pt)
      C  – MON        (blank signing cell; pre-filled 'C' if camper absent that day)
      D  – TUES
      E  – WED
      F  – THURS
      G  – FRI
      H  – Enrolled
    """
    from openpyxl.worksheet.pagebreak import Break

    if report_date is None:
        report_date = date.today()

    campers_sorted = sorted(
        campers,
        key=lambda c: (_bunk_sort_key(c["bunk"]), c["name"])
    )

    seen, groups = [], {}
    for c in campers_sorted:
        bk = c["bunk"]
        if bk not in groups:
            groups[bk] = []
            seen.append(bk)
        groups[bk].append(c)

    # ---- Styles ----
    _thin = Side(style="thin")
    _med  = Side(style="medium")
    T_ALL = Border(left=_thin, right=_thin, top=_thin, bottom=_thin)

    F_WH_LG   = Font(name="Calibri", bold=True,  size=16, color=WHITE)
    F_WH_SM   = Font(name="Calibri", bold=True,  size=11, color=WHITE)
    F_WH_DAY  = Font(name="Calibri", bold=True,  size=16, color=WHITE)
    F_LABEL   = Font(name="Calibri", bold=True,  size=22)
    F_NAME    = Font(name="Calibri", bold=True,  size=16)
    F_ENROLL  = Font(name="Calibri", bold=False, size=16)
    F_COUNT   = Font(name="Calibri", bold=True,  size=16)
    F_WEEK_HDR = Font(name="Calibri", bold=True,  size=24)
    F_DATE_HDR = Font(name="Calibri", bold=True,  size=12)
    F_ABSENT   = Font(name="Calibri", bold=False, size=20, color="999999")

    BRAND_FILL = PatternFill("solid", fgColor=BRAND)
    ALT_FILL   = PatternFill("solid", fgColor="D9D9D9")
    CTR        = Alignment(horizontal="center", vertical="center")
    RIGHT_AL   = Alignment(horizontal="right",  vertical="center")
    VERT_CTR   = Alignment(horizontal="center", vertical="center", text_rotation=90)

    # ---- Row 1: Week/date header (pre-filled; user completes week # and dates in Excel) ----
    ws.row_dimensions[1].height = 36
    ws.merge_cells(start_row=1, start_column=2, end_row=1, end_column=7)
    c = ws.cell(row=1, column=2, value=(f"WEEK {week_num}: {_week_dates(week_num)}" if week_num else "WEEK # :"))
    c.font = F_WEEK_HDR; c.alignment = CTR

    date_str = (report_date.strftime("%-m/%-d/%Y") if os.name != "nt"
                else report_date.strftime("%#m/%#d/%Y"))
    c = ws.cell(row=1, column=8, value=f"Printed: {date_str}")
    c.font = Font(name="Calibri", bold=True, size=9)
    c.alignment = Alignment(horizontal="right", vertical="center", wrap_text=True)

    # ---- Row 2: column headers ----
    ws.row_dimensions[2].height = 20
    hdr = [("A", None, ""),
           ("B", F_WH_LG,  "Camper"),
           ("C", F_WH_DAY, "MON"),
           ("D", F_WH_DAY, "TUES"),
           ("E", F_WH_DAY, "WED"),
           ("F", F_WH_DAY, "THURS"),
           ("G", F_WH_DAY, "FRI"),
           ("H", F_WH_SM,  "Enrolled")]
    for col_letter, font, label in hdr:
        col_idx = ord(col_letter) - ord("A") + 1
        c = ws.cell(row=2, column=col_idx, value=label or None)
        if font:
            c.font = font; c.fill = BRAND_FILL; c.alignment = CTR; c.border = T_ALL

    _DAY_LETTERS = "MTWRF"  # maps day index 0-4 to letters used in enrolled string

    def _write_attend_row(r, label, enrolled):
        ws.row_dimensions[r].height = 31.5
        alt = (r % 2 == 0)
        c = ws.cell(row=r, column=2, value=label)
        c.font = F_NAME; c.alignment = CTR; c.border = T_ALL
        if alt: c.fill = ALT_FILL
        # Cols C–G: blank signing cells; pre-fill 'C' for days not attending
        for di, letter in enumerate(_DAY_LETTERS):
            col = 3 + di
            if enrolled and letter not in enrolled:
                c = ws.cell(row=r, column=col, value="C")
                c.font = F_ABSENT; c.alignment = CTR; c.border = T_ALL
            else:
                c = ws.cell(row=r, column=col); c.border = T_ALL
            if alt: c.fill = ALT_FILL
        c = ws.cell(row=r, column=8, value=enrolled or None)
        c.font = F_ENROLL; c.alignment = CTR; c.border = T_ALL
        if alt: c.fill = ALT_FILL

    # FT CITs assigned to a real bunk (via the "CIT Bunk" column) get an extra
    # line on that bunk's page after the count — matched by bunk name (no #).
    # CITs assigned to non-bunks (Arts & Crafts, Athletics, …) match nothing
    # and so appear only on their own FT CITs page.
    bunk_by_stripped = {_label_bunk(bk).lower(): bk for bk in seen}
    cit_by_bunk = {}
    for c in campers_sorted:
        cb = (c.get("cit_bunk") or "").strip()
        if not cb:
            continue
        target = bunk_by_stripped.get(cb.lower())
        if target and target != c["bunk"]:
            cit_by_bunk.setdefault(target, []).append(c)

    def _disp(c):
        # FT CITs show their assigned area (master "CIT Bunk" column) in parens
        area = (c.get("cit_bunk") or "").strip()
        return f"{c['name']} ({area})" if area else c["name"]

    # ---- Data rows: one bunk per page ----
    row = 3
    total_count = 0

    for bk_idx, bk in enumerate(seen):
        group    = groups[bk]
        count    = len(group)
        total_count += count
        bk_start = row

        for camper in group:
            _write_attend_row(row, _disp(camper), camper["enrolled"])
            row += 1

        # Subtotal row (count of assigned campers — CITs below are not counted)
        ws.row_dimensions[row].height = 31.5
        use_alt = (row % 2 == 0)
        c = ws.cell(row=row, column=2, value=count)
        c.font = F_COUNT; c.alignment = CTR; c.border = T_ALL
        if use_alt: c.fill = ALT_FILL
        for ci in range(3, 9):
            c = ws.cell(row=row, column=ci)
            c.border = T_ALL
            if use_alt: c.fill = ALT_FILL
        bk_end = row
        row += 1

        # FT CITs assigned to this bunk — extra line(s) after the count
        for cit in sorted(cit_by_bunk.get(bk, []), key=lambda x: x["name"].lower()):
            _write_attend_row(row, _disp(cit), cit["enrolled"])
            bk_end = row
            row += 1

        # Merge col A for the whole bunk block (incl. CIT lines), rotate 90°
        ws.merge_cells(start_row=bk_start, start_column=1,
                       end_row=bk_end,     end_column=1)
        c = ws.cell(row=bk_start, column=1, value=bk)
        c.font = F_LABEL; c.alignment = VERT_CTR

        # Page break after each bunk (except the last)
        if bk_idx < len(seen) - 1:
            ws.row_breaks.append(Break(id=bk_end))

    # Grand total row
    ws.row_dimensions[row].height = 31.5
    use_alt = (row % 2 == 0)
    c = ws.cell(row=row, column=2, value=total_count)
    c.font = F_COUNT; c.alignment = CTR; c.border = T_ALL
    if use_alt: c.fill = ALT_FILL
    for ci in range(3, 9):
        ws.cell(row=row, column=ci).border = T_ALL

    # ---- Column widths ----
    ws.column_dimensions["A"].width = 4
    ws.column_dimensions["B"].width = 32
    for col in ["C", "D", "E", "F", "G"]:
        ws.column_dimensions[col].width = 12
    ws.column_dimensions["H"].width = 10

    # ---- Footer: legend printed on every page (spread across L/C/R) ----
    ws.oddFooter.left.text   = "&32✓&16 = Camper in Attendance"
    ws.oddFooter.center.text = "&32C&16 = Camper Confirmed Absent"
    ws.oddFooter.right.text  = "&32O&16 = Camper Not Present"

    # ---- Print settings ----
    ws.print_title_rows = "1:2"          # repeat week header + column headers on every page
    ws.page_setup.orientation = "portrait"
    ws.page_setup.fitToPage   = True
    ws.page_setup.fitToWidth  = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True

    # ---- Margins (inches) ----
    # Larger bottom margin so the big 32pt legend footer never overlaps the
    # last data row on multi-page reports.
    ws.page_margins.top    = 0.5
    ws.page_margins.bottom = 0.9
    ws.page_margins.left   = 0.25
    ws.page_margins.right  = 0.25
    ws.page_margins.footer = 0.3


# ---------------------------------------------------------------------------
# AM / PM Extend parser + builder
# ---------------------------------------------------------------------------

_EXT_TIME_RE    = re.compile(r"(?:Hours|Drop-off)\s+(\d+(?::\d+)?)\s*[-–]", re.IGNORECASE)
_PM_EXT_TIME_RE = re.compile(r"Pick-up\s+\d+(?::\d+)?\s*[^\d\s]\s*(\d+(?::\d+)?)", re.IGNORECASE)

# Week date ranges for the 2026 camp season (shown in report headers)
_WEEK_DATES = [
    "June 22 – 26",
    "June 29 – July 3",
    "July 6 – July 10",
    "July 13 – July 17",
    "July 20 – July 24",
    "July 27 – July 31",
    "August 3 – August 7",
    "August 10 – August 14",
]


def _week_dates(week_num) -> str:
    """Date-range text for a camp week (1-8); empty if out of range."""
    return _WEEK_DATES[week_num - 1] if (week_num and 1 <= week_num <= 8) else ""


def _parse_ext_time(token: str) -> datetime.time:
    """Convert '7', '7:30', '8', '8:30' to datetime.time."""
    if ":" in token:
        h, m = token.split(":")
        return datetime.time(int(h), int(m))
    return datetime.time(int(token), 0)


def parse_extend(file_bytes: bytes, period: str = "am") -> list:
    """
    Parse raw AM/PM Extended Hours export (XLSX or CSV).

    Expected columns (0-indexed):
      0  row#
      1  Last name
      2  First name
      3  Bunk name
      4  Enrollment string  (e.g. "AM Extended Hours 8-8:30 drop-off: 5 Days 6 Wks")
      5  Monday?   (Yes / No / blank)
      6  Tuesday?
      7  Wednesday?
      8  Thursday?
      9  Friday?
    """
    if file_bytes[:4] == b'PK\x03\x04':
        from openpyxl import load_workbook as _lw
        _wb = _lw(filename=io.BytesIO(file_bytes), read_only=True, data_only=True)
        _ws = _wb.active
        rows = [[str(c.value) if c.value is not None else "" for c in r]
                for r in _ws.iter_rows()]
        _wb.close()
    else:
        content = file_bytes.decode("utf-8-sig", errors="replace")
        rows = list(csv.reader(io.StringIO(content)))

    if period == "am":
        # "AM Extended Hours drop-off" (legacy) or bare "Drop-off" (current)
        keywords = ["am extended", "drop-off"]
    else:
        # "PM Extended Hours Pick-up" (legacy) or bare "Pick-up" (current)
        keywords = ["pm extended", "pick-up"]
    campers = []
    for row in rows[1:]:
        if len(row) < 4 or not str(row[0]).strip().isdigit():
            continue
        enrollment = str(row[4]).strip() if len(row) > 4 else ""
        enroll_lower = enrollment.lower()
        if not any(kw in enroll_lower for kw in keywords):
            continue

        last  = str(row[1]).strip()
        first = str(row[2]).strip()
        bunk  = str(row[3]).strip()
        mon    = str(row[5]).strip() if len(row) > 5 else ""
        tue    = str(row[6]).strip() if len(row) > 6 else ""
        wed    = str(row[7]).strip() if len(row) > 7 else ""
        thu    = str(row[8]).strip() if len(row) > 8 else ""
        fri    = str(row[9]).strip() if len(row) > 9 else ""
        gender = str(row[10]).strip() if len(row) > 10 else ""

        # Extract time from enrollment string
        # AM: use start time (before dash); PM: use end/pickup time (after dash)
        time_re = _PM_EXT_TIME_RE if period == "pm" else _EXT_TIME_RE
        m = time_re.search(enrollment)
        start_time = _parse_ext_time(m.group(1)) if m else None

        # Days/Wk
        # days_sched: full set of scheduled day letters (always populated).
        # days_wk:    compact label, shown only for a partial week (blank when all 5).
        any_specified = any(d.lower() in ("yes", "no") for d in [mon, tue, wed, thu, fri])
        if any_specified:
            days_sched = (
                ("M" if mon.lower() == "yes" else "") +
                ("T" if tue.lower() == "yes" else "") +
                ("W" if wed.lower() == "yes" else "") +
                ("R" if thu.lower() == "yes" else "") +
                ("F" if fri.lower() == "yes" else "")
            )
        else:
            days_sched = "MTWRF"
        days_wk = "" if days_sched == "MTWRF" else days_sched

        campers.append({
            "name":       f"{last}, {first}",
            "bunk":       bunk,
            "time":       start_time,
            "days_wk":    days_wk,
            "days_sched": days_sched,
            "gender":     gender,
        })

    # Sort alphabetically by name
    campers.sort(key=lambda c: c["name"].lower())
    return campers


def build_extend_sheet(ws, campers: list, period: str, week_num: int = None) -> None:
    """
    Build the single sheet for AM/PM Extend report.

    AM layout (9 cols A-I):  CAMPER, BUNK, TIME, MON-FRI (1 col each), Days/Wk
    PM layout (14 cols A-N): CAMPER, BUNK, TIME, MON-FRI (2 merged cols each), Days/Wk
    """
    _thin = Side(style="thin")
    _med  = Side(style="medium")
    T_BOT_THIN = Border(bottom=_thin)
    T_BOT_MED  = Border(bottom=_med)
    T_ALL_THIN = Border(left=_thin, right=_thin, top=_thin, bottom=_thin)
    # Full medium box around header cells (split for merged PM day pairs)
    B_HDR_FULL = Border(left=_med, right=_med, top=_med, bottom=_med)
    B_HDR_L    = Border(left=_med, top=_med, bottom=_med)
    B_HDR_R    = Border(right=_med, top=_med, bottom=_med)

    if period == "pm":
        HDR_COLOR = "6A1330"
        ALT_COLOR = "DCDCDC"
        FONT_NAME = "Aptos Narrow"
        DAYS_COL  = 14
        SIGN_RANGE = range(4, 14)   # D–M (10 signing cols, 2 per day)
    else:
        HDR_COLOR = BRAND
        ALT_COLOR = "D9D9D9"
        FONT_NAME = "Calibri"
        DAYS_COL  = 9
        SIGN_RANGE = range(4, 9)    # D–H (5 signing cols)

    HDR_FILL = PatternFill("solid", fgColor=HDR_COLOR)
    ALT_FILL = PatternFill("solid", fgColor=ALT_COLOR)

    # PM Extend gets the enhanced layout (bigger names/rows, header boxes,
    # day X marks, footer). AM Extend keeps its original appearance.
    is_pm = (period == "pm")

    F_HDR  = Font(name=FONT_NAME, bold=True,  size=11, color=WHITE)
    F_NAME = Font(name=FONT_NAME, bold=False, size=14 if is_pm else 11)   # larger PM names
    F_BUNK = Font(name=FONT_NAME, bold=False, size=9)
    F_TIME = Font(name=FONT_NAME, bold=True,  size=11)
    F_DAYS = Font(name=FONT_NAME, bold=False, size=11)
    F_WEEK = Font(name=FONT_NAME, bold=True,  size=11)
    F_X    = Font(name=FONT_NAME, bold=False, size=14, color="808080")  # greyed X (visible on shaded rows)
    F_DAY  = Font(name=FONT_NAME, bold=True,  size=16, color=WHITE)      # AM: large day names
    F_INSTR= Font(name=FONT_NAME, bold=True,  italic=True, size=12, color=WHITE)

    DATA_H = 30.0 if is_pm else 23.75   # taller PM rows → ~25 names per page
    HDR_BORDER = B_HDR_FULL if is_pm else T_BOT_MED

    CTR  = Alignment(horizontal="center", vertical="center")
    WRAP = Alignment(horizontal="center", vertical="center", wrap_text=True)
    LEFT = Alignment(horizontal="left",   vertical="center")

    # ---- Row 1: WEEK label ----
    ws.row_dimensions[1].height = 14.65
    c = ws.cell(row=1, column=1, value=(f"WEEK {week_num}: {_week_dates(week_num)}" if week_num else "WEEK:"))
    c.font = F_WEEK

    def _hdr(col, val, align=CTR, border=HDR_BORDER):
        c = ws.cell(row=2, column=col, value=val)
        c.font = F_HDR; c.fill = HDR_FILL
        c.alignment = align; c.border = border

    if period == "pm":
        # ---- Two-row header: large day name (row 2) + Time/Initial (row 3) ----
        ws.row_dimensions[2].height = 26
        ws.row_dimensions[3].height = 16
        # CAMPER / BUNK / TIME / Days/Wk span both header rows
        for col, lbl in [(1, "CAMPER"), (2, "BUNK"), (3, "TIME"), (DAYS_COL, "Days/Wk")]:
            ws.merge_cells(start_row=2, start_column=col, end_row=3, end_column=col)
            cc = ws.cell(row=2, column=col, value=lbl)
            cc.font = F_HDR; cc.fill = HDR_FILL; cc.alignment = CTR; cc.border = B_HDR_FULL
            r3 = ws.cell(row=3, column=col); r3.fill = HDR_FILL; r3.border = B_HDR_FULL
        day_pairs = [(4, 5, "MON"), (6, 7, "TUES"), (8, 9, "WED"),
                     (10, 11, "THURS"), (12, 13, "FRI")]
        for c1, c2, lbl in day_pairs:
            # Row 2: large day name across the day's two columns
            ws.merge_cells(start_row=2, start_column=c1, end_row=2, end_column=c2)
            dn = ws.cell(row=2, column=c1, value=lbl)
            dn.font = F_DAY; dn.fill = HDR_FILL; dn.alignment = CTR; dn.border = B_HDR_L
            ws.cell(row=2, column=c2).border = B_HDR_R
            # Row 3: Time | Initial sub-labels
            t1 = ws.cell(row=3, column=c1, value="Time")
            t1.font = F_HDR; t1.fill = HDR_FILL; t1.alignment = CTR; t1.border = B_HDR_L
            t2 = ws.cell(row=3, column=c2, value="Initial")
            t2.font = F_HDR; t2.fill = HDR_FILL; t2.alignment = CTR; t2.border = B_HDR_R
        data_start = 4
    else:
        # ---- AM: two-row header — large day names (row 2) +
        #      a thin "Indicate arrival time each day" instruction (row 3) ----
        ws.row_dimensions[2].height = 24
        ws.row_dimensions[3].height = 16
        # CAMPER / BUNK / TIME / Days/Wk span both header rows
        for col, lbl in [(1, "CAMPER"), (2, "BUNK"), (3, "TIME"), (DAYS_COL, "Days/Wk")]:
            ws.merge_cells(start_row=2, start_column=col, end_row=3, end_column=col)
            cc = ws.cell(row=2, column=col, value=lbl)
            cc.font = F_HDR; cc.fill = HDR_FILL; cc.alignment = CTR
            r3 = ws.cell(row=3, column=col)
            r3.fill = HDR_FILL; r3.border = T_BOT_MED
        # Large day-name headers (row 2), no Date/Time sublabels
        for ci, lbl in [(4, "MON"), (5, "TUES"), (6, "WED"), (7, "THURS"), (8, "FRI")]:
            cc = ws.cell(row=2, column=ci, value=lbl)
            cc.font = F_DAY; cc.fill = HDR_FILL; cc.alignment = CTR
        # Instruction row (row 3) merged across the five day columns
        ws.merge_cells(start_row=3, start_column=4, end_row=3, end_column=8)
        inst = ws.cell(row=3, column=4, value="Indicate arrival time each day")
        inst.font = F_INSTR; inst.fill = HDR_FILL; inst.alignment = CTR
        for ci in range(4, 9):
            ws.cell(row=3, column=ci).fill = HDR_FILL
            ws.cell(row=3, column=ci).border = T_BOT_MED
        data_start = 4

    # ---- Data rows ----
    for i, camper in enumerate(campers):
        r = i + data_start
        ws.row_dimensions[r].height = DATA_H
        af = ALT_FILL if (i % 2 == 1) else None

        def _set(col, val, font, align=CTR):
            cell = ws.cell(row=r, column=col, value=val)
            cell.font = font; cell.alignment = align; cell.border = T_BOT_THIN
            if af: cell.fill = af

        _set(1, camper["name"], F_NAME, LEFT)
        _set(2, camper["bunk"], F_BUNK)

        t = camper["time"]
        time_str = (f"{t.hour}:{t.minute:02d}" if t.minute else str(t.hour)) if t else None
        _set(3, time_str, F_TIME)

        for ci in SIGN_RANGE:
            cell = ws.cell(row=r, column=ci)
            cell.border = T_ALL_THIN
            if af: cell.fill = af

        # Lightly greyed X on days the camper is not scheduled
        sched = camper.get("days_sched", "MTWRF")
        if is_pm:
            day_cells = [(4, 5, "M"), (6, 7, "T"), (8, 9, "W"),
                         (10, 11, "R"), (12, 13, "F")]
            for c1, c2, letter in day_cells:
                if letter not in sched:
                    ws.merge_cells(start_row=r, start_column=c1, end_row=r, end_column=c2)
                    cell = ws.cell(row=r, column=c1, value="X")
                    cell.font = F_X; cell.alignment = CTR
                    if af: cell.fill = af
        else:
            # AM: single column per day (D–H = cols 4–8)
            for di, letter in enumerate("MTWRF"):
                if letter not in sched:
                    cell = ws.cell(row=r, column=4 + di, value="X")
                    cell.font = F_X; cell.alignment = CTR
                    if af: cell.fill = af

        _set(DAYS_COL, camper["days_wk"] or None, F_DAYS)

    # ---- Column widths ----
    ws.column_dimensions["B"].width = 9
    if period == "pm":
        ws.column_dimensions["A"].width = 22   # wider CAMPER — names were cut off
        ws.column_dimensions["C"].width = 6    # skinnier TIME to compensate
        # 10 signing columns (D–M, 2 per day) all the same width
        for col_i in range(4, DAYS_COL):
            ws.column_dimensions[get_column_letter(col_i)].width = 8.43
        ws.column_dimensions[get_column_letter(DAYS_COL)].width = 11.6
    else:
        ws.column_dimensions["A"].width = 18
        ws.column_dimensions["C"].width = 9
        for col in ["D", "E", "F", "G", "H", "I"]:
            ws.column_dimensions[col].width = 11.6

    # ---- Print settings ----
    ws.page_setup.orientation = "portrait"
    ws.page_setup.fitToPage   = True
    ws.page_setup.fitToWidth  = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.print_title_rows = "1:3"

    # ---- Margins + footers ----
    if is_pm:
        # Left margin 0.5" leaves room for hole punches in a binder
        ws.page_margins.top    = 0.5
        ws.page_margins.bottom = 0.5
        ws.page_margins.left   = 0.5
        ws.page_margins.right  = 0.25
        ws.page_margins.header = 0.3
        ws.page_margins.footer = 0.25
        ws.print_options.horizontalCentered = False

        # Footer: page number (center) + printed date (bottom-right, small)
        ws.oddFooter.center.text  = "&12&P of &N"
        ws.evenFooter.center.text = "&12&P of &N"
        ws.oddFooter.right.text   = "&8Printed: &D"
        ws.evenFooter.right.text  = "&8Printed: &D"
    else:
        ws.page_margins.top    = 0.7
        ws.page_margins.bottom = 0.4
        ws.page_margins.left   = 0.25
        ws.page_margins.right  = 0.25
        ws.page_margins.header = 0.3
        ws.page_margins.footer = 0.2
        ws.print_options.horizontalCentered = True

        # Small page number (center) + printed date (bottom-right), same size
        ws.oddFooter.center.text  = "&8&P of &N"
        ws.evenFooter.center.text = "&8&P of &N"
        ws.oddFooter.right.text   = "&8Printed: &D"
        ws.evenFooter.right.text  = "&8Printed: &D"


# ---------------------------------------------------------------------------
# PM GRP Extend helpers
# ---------------------------------------------------------------------------

def assign_pm_groups(campers: list, config: dict) -> list:
    """
    Annotate each (PM-extended) camper with their group code and return them
    sorted by group order, bunk number, then name.

    Grp is resolved from bunk_config.json (by bunk number), except bunks
    30 & 31 (Part-time CITs) which split by gender: Girls → Up1, Boys → Up2.
    If a camper has no gender (e.g. from a master sheet without that column),
    fall back to the Maroon/Silver label in the bunk name.
    """
    # Build number → grp, name → grp, AND track the minimum bunk number per group
    num_to_grp:  dict = {}
    name_to_grp: dict = {}   # fallback for bunks with no leading number
    grp_min_bunk: dict = {}

    for camp in config.get("camps", []):
        for bunk in camp.get("bunks", []):
            grp  = bunk.get("grp", "").strip()
            num  = bunk.get("number")
            name = bunk.get("name", "").strip().lower()
            if grp and num is not None:
                num_to_grp[num] = grp
                if grp not in grp_min_bunk or num < grp_min_bunk[grp]:
                    grp_min_bunk[grp] = num
            if grp and name:
                name_to_grp[name] = grp

    # Dynamic group order — sorted by the first bunk number in each group
    grp_idx = {g: i for i, g in enumerate(
        sorted(grp_min_bunk.keys(), key=lambda g: grp_min_bunk[g])
    )}

    # Bunks whose group depends on gender rather than bunk_config
    _GENDER_SPLIT_BUNKS = {30, 31}

    for c in campers:
        m = re.match(r'^(\d+)', c["bunk"].strip())
        bunk_num = int(m.group(1)) if m else None
        c["bunk_num"] = bunk_num if bunk_num is not None else 999
        if bunk_num in _GENDER_SPLIT_BUNKS:
            g = c.get("gender", "").lower()
            bl = c["bunk"].lower()
            if "girl" in g or "female" in g or g == "f":
                c["grp"] = "Up1"
            elif "boy" in g or "male" in g or g == "m":
                c["grp"] = "Up2"
            elif "maroon" in bl:      # gender missing — fall back to bunk label
                c["grp"] = "Up1"
            elif "silver" in bl:
                c["grp"] = "Up2"
            else:
                c["grp"] = "Up2"
        elif bunk_num is not None:
            # Numeric bunk — look up by number
            c["grp"] = num_to_grp.get(bunk_num, "Unknown")
        else:
            # No leading number (e.g. "FT CITs") — look up by name
            c["grp"] = name_to_grp.get(c["bunk"].strip().lower(), "Unknown")

    campers.sort(key=lambda c: (grp_idx.get(c["grp"], 99), c["bunk_num"], c["name"].lower()))
    return campers


def parse_pm_grp_extend(file_bytes: bytes, config: dict) -> list:
    """Parse PM Extended data and annotate each camper with their group code."""
    campers = parse_extend(file_bytes, period="pm")
    return assign_pm_groups(campers, config)


# ---------------------------------------------------------------------------
# Master sheet parser — one upload that every report can derive from
# ---------------------------------------------------------------------------

def _read_rows(file_bytes: bytes) -> list:
    """Read CSV or XLSX bytes into a list of string rows."""
    if file_bytes[:4] == b'PK\x03\x04':
        from openpyxl import load_workbook as _lw
        _wb = _lw(filename=io.BytesIO(file_bytes), read_only=True, data_only=True)
        _ws = _wb.active
        rows = [[str(c.value) if c.value is not None else "" for c in r]
                for r in _ws.iter_rows()]
        _wb.close()
        return rows
    content = file_bytes.decode("utf-8-sig", errors="replace")
    return list(csv.reader(io.StringIO(content)))


def _looks_like_master(header: list) -> bool:
    """Identify the master sheet by its distinctive column headers."""
    hl = [str(h).lower() for h in header]
    joined = " | ".join(hl)
    return ("enrollment extra" in joined
            and any("session" in h for h in hl)
            and any("bunk" in h for h in hl))


def is_master(file_bytes: bytes) -> bool:
    """Quick check (header only) of whether a file is a master sheet."""
    try:
        rows = _read_rows(file_bytes)
        return bool(rows) and _looks_like_master(rows[0])
    except Exception:
        return False


def parse_master(file_bytes: bytes):
    """
    Parse the camp 'master' export (one row per camper) into rich records that
    every report can derive from. Returns None if the file isn't a master.

    Columns are matched by header name (year prefixes like '2026 >' are
    ignored), so column order and extra columns don't matter.
    """
    rows = _read_rows(file_bytes)
    if not rows or not _looks_like_master(rows[0]):
        return None

    header      = rows[0]
    last_col    = _detect_col(header, ["last"],          1)
    first_col   = _detect_col(header, ["first"],         2)
    bunk_col    = _detect_col(header, ["bunk"],          3)
    session_col = _detect_col(header, ["session"],       4)
    age_col     = _detect_col(header, ["age"],           5)
    grade_col   = _detect_col(header, ["grade"],         6)
    mon_col     = _detect_col(header, ["monday"],        7)
    tue_col     = _detect_col(header, ["tuesday"],       8)
    wed_col     = _detect_col(header, ["wednesday"],     9)
    thu_col     = _detect_col(header, ["thursday"],      10)
    fri_col     = _detect_col(header, ["friday"],        11)
    extra_col   = _detect_col(header, ["enrollment", "extra"], 12)
    driver_col  = _detect_col(header, ["driver"],        13)
    stop_col    = _detect_col(header, ["stop"],          None)
    gender_col  = _detect_col(header, ["gender"],        None)
    if gender_col is None:
        gender_col = _detect_col(header, ["sex"],        None)
    cit_bunk_col = _detect_col(header, ["cit", "bunk"],  None)   # FT CIT's assigned bunk

    def _val(row, col):
        return str(row[col]).strip() if (col is not None and col < len(row)) else ""

    records = []
    seen = set()   # de-duplicate campers listed more than once (same name + bunk)
    for row in rows[1:]:
        if len(row) < 4 or not str(row[0]).strip().isdigit():
            continue

        last  = _val(row, last_col)
        first = _val(row, first_col)
        bunk  = _norm(_val(row, bunk_col))
        dup_key = (f"{last}, {first}".strip().lower(), bunk.lower())
        if dup_key in seen:
            continue   # already have this camper — skip the duplicate row
        seen.add(dup_key)
        sessions  = _val(row, session_col)
        raw_age   = _val(row, age_col)
        grade     = normalize_grade(_val(row, grade_col))
        mon = _val(row, mon_col); tue = _val(row, tue_col); wed = _val(row, wed_col)
        thu = _val(row, thu_col); fri = _val(row, fri_col)
        extra   = _val(row, extra_col)
        raw_drv = _val(row, driver_col)
        driver  = "" if raw_drv.lower() in ("", "none", "nan", "n/a", "#n/a") else raw_drv
        gender  = _val(row, gender_col)

        try:
            stop_val = int(float(_val(row, stop_col)))
        except (ValueError, TypeError):
            stop_val = None
        try:
            age_val = float(raw_age)
        except (ValueError, TypeError):
            age_val = raw_age if raw_age else None

        # Weeks 1-8 from the session text
        weeks = [0] * 8
        for part in sessions.split(","):
            m = WEEK_RE.search(part)
            if m:
                wk = int(m.group(1))
                if 1 <= wk <= 8:
                    weeks[wk - 1] = 1

        # Day schedule from Mon-Fri columns (all blank → attends all 5 days)
        any_day = any(d.lower() in ("yes", "no") for d in [mon, tue, wed, thu, fri])
        if any_day:
            day_m = "M" if mon.lower() == "yes" else None
            day_t = "T" if tue.lower() == "yes" else None
            day_w = "W" if wed.lower() == "yes" else None
            day_r = "R" if thu.lower() == "yes" else None
            day_f = "F" if fri.lower() == "yes" else None
            days_sched = "".join(x for x in [day_m, day_t, day_w, day_r, day_f] if x)
        else:
            day_m, day_t, day_w, day_r, day_f = "M", "T", "W", "R", "F"
            days_sched = "MTWRF"
        partial = "" if days_sched == "MTWRF" else days_sched

        # Extended hours from the 'Enrollment extra names' column:
        #   "Drop-off ... AM" → AM drop-off time;  "Pick-up ... PM" → PM pickup time
        am_m = _EXT_TIME_RE.search(extra)
        pm_m = _PM_EXT_TIME_RE.search(extra)
        am_time = _parse_ext_time(am_m.group(1)) if am_m else None
        pm_time = _parse_ext_time(pm_m.group(1)) if pm_m else None

        records.append({
            "name":       f"{last}, {first}",
            "bunk":       bunk,
            "weeks":      weeks,
            "days":       [day_m, day_t, day_w, day_r, day_f],
            "enrolled":   partial,         # Group Attendance
            "days_sched": days_sched,      # Extend C/X marks
            "days_wk":    partial,         # Extend compact label
            "age":        age_val,
            "grade":      grade,
            "driver":     driver,
            "stop":       stop_val,
            "gender":     gender,
            "am_time":    am_time,
            "pm_time":    pm_time,
            "extra":      extra,           # raw 'Enrollment extra names' text
            "cit_bunk":   _val(row, cit_bunk_col),  # FT CIT's assigned bunk (name only)
        })

    return records


def master_extend_campers(records: list, period: str) -> list:
    """From master records, return campers enrolled in AM (drop-off) or PM
    (pick-up) extended hours, with their pickup/dropoff time, sorted by name."""
    key = "am_time" if period == "am" else "pm_time"
    out = [{**r, "time": r[key]} for r in records if r.get(key) is not None]
    out.sort(key=lambda c: c["name"].lower())
    return out


# ---------------------------------------------------------------------------
# Word label sheet (Avery 5960) for a camp/group
# ---------------------------------------------------------------------------

_DAY_FULL = {"M": "Mon", "T": "Tue", "W": "Wed", "R": "Thu", "F": "Fri"}


def _label_bunk(bunk: str) -> str:
    """Bunk name for labels with the leading number removed
    (e.g. '01 Munchkins' -> 'Munchkins', '30 PT CITs - Maroon' -> 'PT CITs - Maroon')."""
    return re.sub(r"^\s*\d+\s+", "", str(bunk)).strip()


def _label_days_text(sched: str) -> str:
    """Readable 'days attending' text from a schedule string like 'MWF'.
    Full-week (Mon-Fri) campers get a blank line."""
    sched = sched or "MTWRF"
    if sched == "MTWRF":
        return ""   # attends all week — leave blank
    return ", ".join(_DAY_FULL[c] for c in "MTWRF" if c in sched)


def _avery5960_docx(rows3: list) -> bytes:
    """Render a list of (line1, line2, line3) tuples onto an Avery 5960 sheet
    (3 x 10, 2.625" x 1"). Returns the .docx bytes."""
    import io as _io
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Inches(8.5)
    sec.page_height   = Inches(11)
    sec.top_margin    = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin   = Inches(0.19)
    sec.right_margin  = Inches(0.19)

    COLS = 3
    n = len(rows3)
    nrows = max(1, (n + COLS - 1) // COLS)
    col_w = [Inches(2.625), Inches(0.125), Inches(2.625), Inches(0.125), Inches(2.625)]
    table = doc.add_table(rows=nrows, cols=5)
    table.allow_autofit = False
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)

    def _fill(cell, l1, l2, l3):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for i, (text, sz, bold) in enumerate([(l1, 10, True), (l2, 14, True), (l3, 10, False)]):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0
            run = p.add_run(text or "")
            run.font.size = Pt(sz); run.font.bold = bold

    li = 0
    for ri in range(nrows):
        row = table.rows[ri]
        trPr = row._tr.get_or_add_trPr()
        trH = OxmlElement("w:trHeight"); trH.set(qn("w:val"), "1440"); trH.set(qn("w:hRule"), "exact")
        trPr.append(trH)
        for ci in range(5):
            cell = row.cells[ci]
            cell.width = col_w[ci]
            if ci % 2 == 1:
                continue   # spacer column
            if li < n:
                _fill(cell, *rows3[li]); li += 1

    buf = _io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()


def _group_bunks(config: dict, group_name: str):
    """Return a predicate: does this bunk belong to the named camp/group?"""
    bunk_camp = {}
    for camp in config.get("camps", []):
        cn = camp.get("name", "")
        for b in camp.get("bunks", []):
            bunk_camp[_norm(b.get("name", "")).lower()] = cn
    gl = group_name.strip().lower()

    def _in_group(bk):
        cn = bunk_camp.get(_norm(bk).lower())
        return bool(cn) and cn.strip().lower().startswith(gl)
    return _in_group


def build_group_labels_docx(records: list, config: dict, group_name: str = "Inter") -> tuple:
    """
    Avery 5960 labels for all campers in a camp/group.
    Each label: bunk / camper / days attending. Returns (docx_bytes, count).
    """
    in_group = _group_bunks(config, group_name)
    labels = sorted((r for r in records if in_group(r["bunk"])),
                    key=lambda r: (_bunk_sort_key(r["bunk"]), r["name"].lower()))
    rows3 = [(_label_bunk(r["bunk"]), r["name"], _label_days_text(r.get("days_sched"))) for r in labels]
    return _avery5960_docx(rows3), len(rows3)


def build_jr_transport_labels_docx(records: list, config: dict) -> tuple:
    """
    Junior-group transportation labels in ONE Avery 5960 .docx, with each
    group starting on a fresh page (so each set can print on its own label
    color by page range):

      • Trans    — has 2-Way or PM-Only Transportation; line 3 = "Trans - <driver>"
      • Pm ext   — has a Drop-off enrollment;            line 3 = "Pm ext"
      • Car line — none of the above;                    line 3 = "Car line"

    Returns (docx_bytes, {group: count}, {group: (first_page, last_page)}).
    """
    PER_PAGE = 30   # Avery 5960: 3 x 10

    in_group = _group_bunks(config, "Junior")
    juniors = sorted((r for r in records if in_group(r["bunk"])),
                     key=lambda r: (_bunk_sort_key(r["bunk"]), r["name"].lower()))

    trans, pmext, carline = [], [], []
    for r in juniors:
        ex = (r.get("extra") or "").lower()
        bk = _label_bunk(r["bunk"])
        if "2-way" in ex or "pm only trans" in ex:
            drv = (r.get("driver") or "").strip()
            trans.append((bk, r["name"], f"Trans - {drv}" if drv else "Trans"))
        elif "drop-off" in ex:
            pmext.append((bk, r["name"], "Pm ext"))
        else:
            carline.append((bk, r["name"], "Car line"))

    groups = [("Trans", trans), ("Pm ext", pmext), ("Car line", carline)]
    nonempty = [(name, rows) for name, rows in groups if rows]

    flat, pages, page = [], {}, 1
    for gi, (name, rows) in enumerate(nonempty):
        npages = (len(rows) + PER_PAGE - 1) // PER_PAGE
        pages[name] = (page, page + npages - 1)
        page += npages
        flat.extend(rows)
        # pad to a full page so the NEXT group starts at the top of a new page
        if gi < len(nonempty) - 1:
            flat.extend([("", "", "")] * (npages * PER_PAGE - len(rows)))

    counts = {"Trans": len(trans), "Pm ext": len(pmext), "Car line": len(carline)}
    return _avery5960_docx(flat), counts, pages


def build_pm_grp_extend_sheet(ws, campers: list) -> None:
    """
    PM GRP EXTEND: landscape. Each group prints on its own page with the group
    name as a banner at the top, then BUNK | CAMPER | Pick Up | MON-FRI | Days,
    and a camper total under the names.
    """
    _thin = Side(style="thin")
    T_ALL = Border(left=_thin, right=_thin, top=_thin, bottom=_thin)

    HDR_FILL   = PatternFill("solid", fgColor="6A1330")
    ALT_FILL   = PatternFill("solid", fgColor="D9D9D9")   # every-other-row shading
    TOTAL_FILL = PatternFill("solid", fgColor="D9D9D9")
    FONT_NAME  = "Aptos Narrow"

    F_GRP    = Font(name=FONT_NAME, bold=True,  size=20, color="000000")  # group banner
    F_HDR    = Font(name=FONT_NAME, bold=True,  size=13, color=WHITE)
    F_DAYHDR = Font(name=FONT_NAME, bold=True,  size=16, color=WHITE)     # larger day names
    F_BUNK   = Font(name=FONT_NAME, bold=False, size=12)                  # smaller bunk
    F_DATA   = Font(name=FONT_NAME, bold=False, size=14)
    F_ABSENT = Font(name=FONT_NAME, bold=False, size=14, color="999999")  # lightly greyed C
    F_TOTAL  = Font(name=FONT_NAME, bold=True,  size=18)                  # larger total

    CTR  = Alignment(horizontal="center", vertical="center")
    LEFT = Alignment(horizontal="left",   vertical="center")

    COL_BUNK, COL_CAMPER, COL_PICKUP, COL_DAY1, COL_DAYS = 1, 2, 3, 4, 9
    LAST_COL = 9
    _DAY_LETTERS = "MTWRF"
    DAY_LABELS   = ["MON", "TUES", "WED", "THURS", "FRI"]
    BANNER_H, HDR_H, DATA_H = 26, 20, 26

    from openpyxl.worksheet.pagebreak import Break

    def _write_headers(hr):
        ws.row_dimensions[hr].height = HDR_H
        for ci, lbl in [(COL_BUNK, "BUNK"), (COL_CAMPER, "CAMPER"),
                        (COL_PICKUP, "Pick Up"), (COL_DAYS, "Days")]:
            c = ws.cell(row=hr, column=ci, value=lbl)
            c.font = F_HDR; c.fill = HDR_FILL; c.alignment = CTR; c.border = T_ALL
        for di, lbl in enumerate(DAY_LABELS):
            c = ws.cell(row=hr, column=COL_DAY1 + di, value=lbl)
            c.font = F_DAYHDR; c.fill = HDR_FILL; c.alignment = CTR; c.border = T_ALL

    def _write_total(hr, count):
        ws.row_dimensions[hr].height = DATA_H
        for ci in range(1, LAST_COL + 1):
            cell = ws.cell(row=hr, column=ci)
            cell.border = T_ALL; cell.fill = TOTAL_FILL
        tc = ws.cell(row=hr, column=COL_CAMPER, value=f"Total: {count}")
        tc.font = F_TOTAL; tc.alignment = LEFT; tc.fill = TOTAL_FILL; tc.border = T_ALL

    # ---- Data rows, grouped; each group on its own page ----
    r = 1
    current_grp = None
    group_count = 0

    for camper in campers:
        if camper["grp"] != current_grp:
            if current_grp is not None:
                _write_total(r, group_count)
                ws.row_breaks.append(Break(id=r))   # page break after the total
                r += 1
            current_grp = camper["grp"]
            group_count = 0
            # Group banner (large black, top-left of the page)
            ws.row_dimensions[r].height = BANNER_H
            b = ws.cell(row=r, column=COL_BUNK, value=current_grp)
            b.font = F_GRP; b.alignment = LEFT
            r += 1
            _write_headers(r)
            r += 1

        ws.row_dimensions[r].height = DATA_H
        af = ALT_FILL if (group_count % 2 == 1) else None
        t = camper["time"]
        time_str = (f"{t.hour}:{t.minute:02d}" if t.minute else str(t.hour)) if t else None

        for col, val, align, fnt in [
            (COL_BUNK,   camper["bunk"],            CTR,  F_BUNK),
            (COL_CAMPER, camper["name"],            LEFT, F_DATA),
            (COL_PICKUP, time_str,                  CTR,  F_DATA),
            (COL_DAYS,   camper["days_wk"] or None, CTR,  F_DATA),
        ]:
            cell = ws.cell(row=r, column=col, value=val)
            cell.font = fnt; cell.alignment = align; cell.border = T_ALL
            if af: cell.fill = af

        # Day columns (MON-FRI): lightly greyed C on days camper is not attending
        sched = camper.get("days_sched", "MTWRF")
        for di, letter in enumerate(_DAY_LETTERS):
            cell = ws.cell(row=r, column=COL_DAY1 + di)
            cell.border = T_ALL
            if af: cell.fill = af
            if letter not in sched:
                cell.value = "C"; cell.font = F_ABSENT; cell.alignment = CTR

        r += 1
        group_count += 1

    # total for the last group
    if current_grp is not None:
        _write_total(r, group_count)

    # ---- Column widths ----
    ws.column_dimensions["A"].width = 14     # BUNK (smaller/narrower)
    ws.column_dimensions["B"].width = 28     # CAMPER
    ws.column_dimensions["C"].width = 11     # Pick Up
    for col in ["D", "E", "F", "G", "H"]:
        ws.column_dimensions[col].width = 13.5
    ws.column_dimensions["I"].width = 14     # Days

    # ---- Print settings ----
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToPage   = True
    ws.page_setup.fitToWidth  = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.print_title_rows = None

    # ---- Margins: wider bottom so the big 32pt legend footer never overlaps -
    ws.page_margins.top    = 0.5
    ws.page_margins.bottom = 0.9
    ws.page_margins.left   = 0.25
    ws.page_margins.right  = 0.25
    ws.page_margins.header = 0.3
    ws.page_margins.footer = 0.3
    ws.print_options.horizontalCentered = True
    ws.print_options.verticalCentered   = False

    # ---- Footer: ✓/C/O legend printed on every page (spread across L/C/R) ----
    ws.oddFooter.left.text   = "&32✓&16 = Camper in Attendance"
    ws.oddFooter.center.text = "&32C&16 = Camper Confirmed Absent"
    ws.oddFooter.right.text  = "&32O&16 = Camper Not Present"


# ---------------------------------------------------------------------------
# Driver Totals builder
# ---------------------------------------------------------------------------

def build_driver_totals_sheet(ws, campers: list, report_date: date, week_num: int = None) -> None:
    """
    Build the Driver Totals sheet.

    Column layout (18 cols A-R):
      A  Child       B  Bunk
      C-J  #1-#8 (week indicators)
      K-O  M T W R F (day letters)
      P  Age   Q  Grade   R  Driver

    Per driver group: data rows (alternating gray/white) → SUM row → COUNT row
    Grand totals: GRAND COUNT row → GRAND SUM row
    Each driver group prints on its own page.
    """
    from openpyxl.worksheet.pagebreak import Break

    F_DATA   = Font(name="Calibri", bold=False, size=16)
    F_BOLD   = Font(name="Calibri", bold=True,  size=16)
    F_HDR    = Font(name="Calibri", bold=True,  size=14)
    F_BANNER = Font(name="Calibri", bold=True,  size=22, color="000000")  # driver name
    F_LEG    = Font(name="Calibri", bold=False, size=11)                  # legend key
    CTR  = Alignment(horizontal="center", vertical="center")
    LEFT = Alignment(horizontal="left",   vertical="center")
    _thin  = Side(style="thin")
    BORDER = Border(left=_thin, right=_thin, top=_thin, bottom=_thin)
    HDR_FILL  = PatternFill("solid", fgColor="D9D9D9")
    ALT_FILL  = PatternFill("solid", fgColor="EEEEEE")  # alternating rows
    AGE_WARN  = PatternFill("solid", fgColor="92D050")  # green: age < 8
    BUNK_WARN = PatternFill("solid", fgColor="FFC000")  # orange: bunks 1-7
    WEEK_FILL = PatternFill("solid", fgColor="FFFF00")  # yellow: selected week

    # Column layout (18 cols A-R; Driver column removed — driver is a banner):
    #   A=Child  B=Stp#  C=Bunk  D-K(4-11)=#1-#8  L-P(12-16)=M/T/W/R/F  Q=Age  R=Grade
    COL_CHILD, COL_STOP, COL_BUNK = 1, 2, 3
    COL_WK1  = 4
    COL_DAY1 = 12
    COL_AGE  = 17
    COL_GRADE = 18
    LAST_COL = 18
    HEADERS = ["Child", "Stp#", "Bunk", "#1", "#2", "#3", "#4", "#5", "#6", "#7", "#8",
               "M", "T", "W", "R", "F", "Age", "Grade"]

    def _c(r, col, val, font=F_DATA, align=CTR, fill=None):
        cell = ws.cell(row=r, column=col, value=val)
        cell.font = font; cell.alignment = align
        if fill: cell.fill = fill
        return cell

    banner_rows = []
    skip_border_rows = set()   # banner + legend rows excluded from separators

    def _headers(hr):
        ws.row_dimensions[hr].height = 22
        for ci, h in enumerate(HEADERS, start=1):
            wk_hl = bool(week_num and ci == COL_WK1 + week_num - 1)
            _c(hr, ci, h, font=F_HDR, align=(LEFT if ci == COL_CHILD else CTR),
               fill=(WEEK_FILL if wk_hl else None))

    def _legend(r):
        """Color key footer under the Age/Grade columns. Returns next free row."""
        wrap = Alignment(horizontal="center", vertical="center", wrap_text=True)
        ws.row_dimensions[r].height = 22
        ws.merge_cells(start_row=r, start_column=COL_DAY1, end_row=r, end_column=COL_GRADE)
        g = ws.cell(row=r, column=COL_DAY1, value="Booster Required")
        g.font = F_LEG; g.alignment = wrap; g.fill = AGE_WARN
        ws.row_dimensions[r + 1].height = 34
        ws.merge_cells(start_row=r + 1, start_column=COL_DAY1, end_row=r + 1, end_column=COL_GRADE)
        o = ws.cell(row=r + 1, column=COL_DAY1,
                    value="Must Walk Camper to Jr Pavilion/Preschool Building in AM")
        o.font = F_LEG; o.alignment = wrap; o.fill = BUNK_WARN
        skip_border_rows.update({r, r + 1})
        return r + 2

    # ----- Group + sort campers by driver -----------------------------------
    driver_groups: dict = {}
    for camper in campers:
        driver_groups.setdefault(camper["driver"] or "(No Driver)", []).append(camper)
    has_stops = any(c.get("stop") is not None for c in campers)
    for drv in driver_groups:
        driver_groups[drv].sort(
            key=(lambda x: (x["stop"] is None, x["stop"] or 0, x["name"].lower()))
            if has_stops else (lambda x: x["name"].lower()))
    sorted_drivers = sorted(driver_groups.keys())

    row = 1
    grand_week_sums = [0] * 8
    grand_count = 0

    for drv in sorted_drivers:
        group = driver_groups[drv]
        week_sums = [0] * 8
        count = len(group)
        grand_count += count

        # --- Driver banner (top of the page) ---
        ws.row_dimensions[row].height = 28
        bn = ws.cell(row=row, column=COL_CHILD, value=drv)
        bn.font = F_BANNER; bn.alignment = LEFT
        banner_rows.append(row)
        row += 1
        _headers(row); row += 1

        for i, camper in enumerate(group):
            ws.row_dimensions[row].height = 22
            base = ALT_FILL if (i % 2 == 1) else None

            bunk_val = camper["bunk"]
            mm = re.match(r"\d+", str(bunk_val or ""))
            bnum = int(mm.group()) if mm else None
            bunk_fill = BUNK_WARN if (bnum is not None and 1 <= bnum <= 7) else base
            attending = bool(week_num and 1 <= week_num <= len(camper["weeks"])
                             and camper["weeks"][week_num - 1] == 1)

            _c(row, COL_CHILD, camper["name"], align=LEFT, fill=(WEEK_FILL if attending else base))
            _c(row, COL_STOP,  camper.get("stop"), fill=base)
            _c(row, COL_BUNK,  bunk_val, fill=bunk_fill)
            for wi, wv in enumerate(camper["weeks"]):
                wf = WEEK_FILL if (week_num and wi == week_num - 1 and wv == 1) else base
                _c(row, COL_WK1 + wi, wv, fill=wf)
                week_sums[wi] += wv; grand_week_sums[wi] += wv
            for di, dv in enumerate(camper["days"]):
                _c(row, COL_DAY1 + di, dv, fill=base)
            age_val = camper["age"]
            agew = AGE_WARN if (isinstance(age_val, (int, float)) and age_val < 8) else base
            _c(row, COL_AGE, age_val, fill=agew)
            _c(row, COL_GRADE, camper["grade"] or None, fill=base)
            row += 1

        # --- Week totals + camper count ---
        ws.row_dimensions[row].height = 22
        _c(row, COL_CHILD, "Week Totals", font=F_BOLD, align=LEFT)
        for wi, wsum in enumerate(week_sums):
            _c(row, COL_WK1 + wi, wsum, font=F_BOLD)
        row += 1
        ws.row_dimensions[row].height = 22
        _c(row, COL_CHILD, f"Total Campers: {count}", font=F_BOLD, align=LEFT)
        row += 1

        # Color-key legend (footer) on each driver's page
        row = _legend(row)

        ws.row_breaks.append(Break(id=row - 1))   # each driver on its own page

    # ----- Grand totals (own page) ------------------------------------------
    ws.row_dimensions[row].height = 28
    gb = ws.cell(row=row, column=COL_CHILD, value="Grand Total — All Drivers")
    gb.font = F_BANNER; gb.alignment = LEFT
    banner_rows.append(row)
    row += 1
    _headers(row); row += 1
    _c(row, COL_CHILD, "Week Totals", font=F_BOLD, align=LEFT)
    for wi, gs in enumerate(grand_week_sums):
        _c(row, COL_WK1 + wi, gs, font=F_BOLD)
    row += 1
    _c(row, COL_CHILD, f"Total Campers: {grand_count}", font=F_BOLD, align=LEFT)

    # ----- Vertical separators (no full grid): right of Stp#, #8, and F -----
    _vert = Side(style="thin", color="000000")
    skip = set(banner_rows) | skip_border_rows
    for r in range(1, row + 1):
        if r in skip:
            continue
        for col in (COL_STOP, COL_BUNK, COL_WK1 + 7, COL_DAY1 + 4):
            ws.cell(row=r, column=col).border = Border(right=_vert)

    # ----- Column widths (fill the landscape width) -------------------------
    ws.column_dimensions["A"].width = 26   # Child
    ws.column_dimensions["B"].width = 6    # Stp#
    ws.column_dimensions["C"].width = 22   # Bunk
    for wi in range(8):
        ws.column_dimensions[get_column_letter(COL_WK1 + wi)].width = 7
    for di in range(5):
        ws.column_dimensions[get_column_letter(COL_DAY1 + di)].width = 3.5
    ws.column_dimensions[get_column_letter(COL_AGE)].width   = 7
    ws.column_dimensions[get_column_letter(COL_GRADE)].width = 9

    # ----- Print settings: landscape, fill width, each driver own page ------
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToPage   = True
    ws.page_setup.fitToWidth  = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.print_title_rows = None
    ws.page_margins.top    = 0.3
    ws.page_margins.bottom = 0.4
    ws.page_margins.left   = 0.25
    ws.page_margins.right  = 0.25
    ws.page_margins.header = 0.2
    ws.page_margins.footer = 0.2

    # Report date in the footer (bottom-right)
    date_str = (report_date.strftime("%-m/%-d/%Y") if os.name != "nt"
                else report_date.strftime("%#m/%#d/%Y"))
    ws.oddFooter.right.text  = f"&12Report Date: {date_str}"
    ws.evenFooter.right.text = f"&12Report Date: {date_str}"


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def process_report(file_bytes: bytes, report_type: str,
                   config: dict, job_id: str, output_dir: str,
                   week_num: int = None) -> dict:

    supported = ("bunk_snapshot", "group_attendance", "am_extend", "pm_extend", "pm_grp_extend", "driver_totals", "inter_labels", "jr_transport_labels")
    if report_type not in supported:
        return {
            "success": False,
            "message": (
                f"Report type '{report_type}' is not configured. "
                f"Currently supported: {', '.join(repr(s) for s in supported)}."
            ),
        }

    report_date = date.today()
    os.makedirs(output_dir, exist_ok=True)

    # Auto-detect a master sheet. When present, every report is derived from it;
    # otherwise each report falls back to its original per-report parser.
    try:
        master = parse_master(file_bytes)
    except Exception:
        master = None

    def _week_filter(campers):
        """For week-specific reports off the master, keep only campers enrolled
        in the selected week. (Old per-report exports carry no week info, so the
        filter only applies when running from a master.)"""
        if master is not None and week_num and 1 <= week_num <= 8:
            return [c for c in campers
                    if c.get("weeks") and len(c["weeks"]) >= week_num
                    and c["weeks"][week_num - 1]]
        return campers

    # ---- Bunk Snapshot ----
    if report_type == "bunk_snapshot":
        try:
            campers = master if master is not None else parse_raw_csv(file_bytes)
        except Exception as e:
            return {"success": False, "message": f"Could not parse file: {e}"}
        if not campers:
            return {"success": False, "message": "No camper data found in file. Check the file format."}

        bunk_lookup   = get_bunk_lookup(config)
        ordered_bunks = get_ordered_bunks(config)

        wb = Workbook()
        ws_report = wb.active
        ws_report.title = "Report"
        ws_totals = wb.create_sheet("Totals")
        build_report_sheet(ws_report, campers, bunk_lookup, ordered_bunks, report_date)
        build_totals_sheet(ws_totals, campers, config, bunk_lookup, report_date)

        out_filename = f"Bunk Snapshot {report_date.strftime('%m%d%Y')}.xlsx"
        out_path = os.path.join(output_dir, out_filename)
        wb.save(out_path)

        return {
            "success":  True,
            "message":  f"Processed {len(campers)} campers successfully.",
            "filename": out_filename,
            "rows":     len(campers),
        }

    # ---- Group Attendance ----
    if report_type == "group_attendance":
        try:
            campers = master if master is not None else parse_group_attendance(file_bytes)
        except Exception as e:
            return {"success": False, "message": f"Could not parse file: {e}"}
        campers = _week_filter(campers)
        if not campers:
            return {"success": False, "message": "No camper data found in file. Check the file format."}

        wb = Workbook()
        ws = wb.active
        ws.title = "Data1"
        build_group_attendance_sheet(ws, campers, config, report_date, week_num=week_num)

        out_filename = f"Group Attendance {report_date.strftime('%m%d%Y')}.xlsx"
        out_path = os.path.join(output_dir, out_filename)
        wb.save(out_path)

        return {
            "success":  True,
            "message":  f"Processed {len(campers)} campers successfully.",
            "filename": out_filename,
            "rows":     len(campers),
        }

    # ---- AM Extend ----
    if report_type == "am_extend":
        try:
            campers = (master_extend_campers(master, "am") if master is not None
                       else parse_extend(file_bytes, period="am"))
        except Exception as e:
            return {"success": False, "message": f"Could not parse file: {e}"}
        campers = _week_filter(campers)
        if not campers:
            return {"success": False, "message": "No AM Extended campers found in file."}

        wb = Workbook()
        ws = wb.active
        ws.title = "AM Extend"
        build_extend_sheet(ws, campers, period="am", week_num=week_num)

        out_filename = f"AM Extend {report_date.strftime('%m%d%Y')}.xlsx"
        out_path = os.path.join(output_dir, out_filename)
        wb.save(out_path)

        return {
            "success":  True,
            "message":  f"Processed {len(campers)} campers successfully.",
            "filename": out_filename,
            "rows":     len(campers),
        }

    # ---- PM Extend ----
    if report_type == "pm_extend":
        try:
            campers = (master_extend_campers(master, "pm") if master is not None
                       else parse_extend(file_bytes, period="pm"))
        except Exception as e:
            return {"success": False, "message": f"Could not parse file: {e}"}
        campers = _week_filter(campers)
        if not campers:
            return {"success": False, "message": "No PM Extended campers found in file."}

        wb = Workbook()
        ws = wb.active
        ws.title = "PM Extend"
        build_extend_sheet(ws, campers, period="pm", week_num=week_num)

        out_filename = f"PM Extend {report_date.strftime('%m%d%Y')}.xlsx"
        out_path = os.path.join(output_dir, out_filename)
        wb.save(out_path)

        return {
            "success":  True,
            "message":  f"Processed {len(campers)} campers successfully.",
            "filename": out_filename,
            "rows":     len(campers),
        }

    # ---- PM GRP Extend ----
    if report_type == "pm_grp_extend":
        try:
            campers = (assign_pm_groups(_week_filter(master_extend_campers(master, "pm")), config)
                       if master is not None
                       else parse_pm_grp_extend(file_bytes, config))
        except Exception as e:
            return {"success": False, "message": f"Could not parse file: {e}"}
        if not campers:
            return {"success": False, "message": "No PM Extended campers found in file."}

        wb = Workbook()
        ws = wb.active
        ws.title = "PM GRP Extend"
        build_pm_grp_extend_sheet(ws, campers)

        out_filename = f"PM GRP Extend {report_date.strftime('%m%d%Y')}.xlsx"
        out_path = os.path.join(output_dir, out_filename)
        wb.save(out_path)

        return {
            "success":  True,
            "message":  f"Processed {len(campers)} campers successfully.",
            "filename": out_filename,
            "rows":     len(campers),
        }

    # ---- Driver Totals ----
    if report_type == "driver_totals":
        try:
            campers = master if master is not None else parse_driver_csv(file_bytes)
        except Exception as e:
            return {"success": False, "message": f"Could not parse file: {e}"}
        # Filter to campers who have a driver assigned
        campers = [c for c in campers if c.get("driver")]
        if not campers:
            return {"success": False, "message": "No campers with a driver assignment found in file."}

        wb = Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        build_driver_totals_sheet(ws, campers, report_date, week_num=week_num)

        out_filename = f"Driver Totals {report_date.strftime('%m%d%Y')}.xlsx"
        out_path = os.path.join(output_dir, out_filename)
        wb.save(out_path)

        return {
            "success":  True,
            "message":  f"Processed {len(campers)} campers across drivers successfully.",
            "filename": out_filename,
            "rows":     len(campers),
        }

    # ---- Inter Labels (Word / Avery 5960) ----
    if report_type == "inter_labels":
        if master is None:
            return {"success": False, "message": "Inter Labels needs a master sheet. Upload the master report."}
        try:
            # Week-dependent: only campers attending the selected week get a label
            docx_bytes, count = build_group_labels_docx(_week_filter(master), config, group_name="Inter")
        except Exception as e:
            return {"success": False, "message": f"Could not build labels: {e}"}
        if not count:
            return {"success": False, "message": "No Group Inter campers attending the selected week."}

        out_filename = f"Inter Labels {report_date.strftime('%m%d%Y')}.docx"
        out_path = os.path.join(output_dir, out_filename)
        with open(out_path, "wb") as f:
            f.write(docx_bytes)

        return {
            "success":  True,
            "message":  f"Created {count} labels for Group Inter.",
            "filename": out_filename,
            "rows":     count,
        }

    # ---- Jr. Transportation Labels (Word / Avery 5960, one file, grouped per page) ----
    if report_type == "jr_transport_labels":
        if master is None:
            return {"success": False, "message": "Jr. Transportation Labels needs a master sheet. Upload the master report."}
        try:
            docx_bytes, counts, pages = build_jr_transport_labels_docx(_week_filter(master), config)
        except Exception as e:
            return {"success": False, "message": f"Could not build labels: {e}"}
        total = sum(counts.values())
        if not total:
            return {"success": False, "message": "No Junior campers attending the selected week."}

        out_filename = f"Junior Labels {report_date.strftime('%m%d%Y')}.docx"
        out_path = os.path.join(output_dir, out_filename)
        with open(out_path, "wb") as f:
            f.write(docx_bytes)

        def _rng(name):
            pr = pages.get(name)
            if not pr:
                return f"{name} 0"
            a, b = pr
            return f"{name} {counts[name]} (page {a})" if a == b else f"{name} {counts[name]} (pages {a}-{b})"
        ranges = ", ".join(_rng(n) for n in ("Trans", "Pm ext", "Car line") if counts[n])

        return {
            "success":  True,
            "message":  f"Created {total} Junior labels — {ranges}. Each group starts on a new page.",
            "filename": out_filename,
            "rows":     total,
        }
